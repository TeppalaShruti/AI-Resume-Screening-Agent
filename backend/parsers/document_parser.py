"""Document parsing: PDF (PyMuPDF), DOCX (python-docx), TXT/MD (plain)."""

from __future__ import annotations

import io
import os
import re
from dataclasses import dataclass, field

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
MAX_BYTES = 10 * 1024 * 1024  # 10 MB per file


class ParseError(Exception):
    """Raised when a document cannot be parsed. Never aborts a batch."""


@dataclass
class ParsedDocument:
    filename: str
    text: str
    pages: int = 1
    warnings: list[str] = field(default_factory=list)


def _clean(text: str) -> str:
    text = text.replace("\x00", " ").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t\u00a0]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _parse_pdf(data: bytes) -> tuple[str, int]:
    import fitz  # PyMuPDF

    with fitz.open(stream=data, filetype="pdf") as doc:
        parts = [page.get_text("text") for page in doc]
        return "\n".join(parts), doc.page_count


def _parse_docx(data: bytes) -> tuple[str, int]:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts), 1


def _parse_txt(data: bytes) -> tuple[str, int]:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding), 1
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore"), 1


def parse_bytes(filename: str, data: bytes) -> ParsedDocument:
    ext = os.path.splitext(filename)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ParseError(f"Unsupported file type '{ext or 'unknown'}'")
    if not data:
        raise ParseError("File is empty")
    if len(data) > MAX_BYTES:
        raise ParseError("File exceeds the 10 MB limit")

    try:
        if ext == ".pdf":
            text, pages = _parse_pdf(data)
        elif ext == ".docx":
            text, pages = _parse_docx(data)
        else:
            text, pages = _parse_txt(data)
    except ParseError:
        raise
    except Exception as exc:  # pragma: no cover - depends on corrupt input
        raise ParseError(f"Could not read document: {exc}") from exc

    cleaned = _clean(text)
    if len(cleaned) < 50:
        raise ParseError("No extractable text found (document may be a scanned image)")

    warnings: list[str] = []
    if len(cleaned) < 300:
        warnings.append("Very little text extracted; extraction quality may be low.")
    return ParsedDocument(filename=filename, text=cleaned, pages=pages, warnings=warnings)


def parse_document(path: str) -> ParsedDocument:
    with open(path, "rb") as handle:
        return parse_bytes(os.path.basename(path), handle.read())
